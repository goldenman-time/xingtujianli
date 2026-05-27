import io
import pytest
from docx import Document
from docx.shared import Pt, Cm
from fastapi.testclient import TestClient

from main import app, html_to_word, ExportWordRequest

client = TestClient(app)


class TestHtmlToWord:
    def test_basic_heading(self):
        html = "<h1>张明的简历</h1><h2>教育背景</h2><h3>详细信息</h3>"
        buffer = html_to_word(html, "张明", "template_15")
        doc = Document(buffer)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "张明的简历" in texts
        assert "教育背景" in texts
        assert "详细信息" in texts

    def test_paragraph(self):
        html = "<p>这是一段普通文字</p>"
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "这是一段普通文字" in texts

    def test_bold_text(self):
        html = "<p><strong>加粗文字</strong>和普通文字</p>"
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        has_bold = False
        for p in doc.paragraphs:
            for run in p.runs:
                if run.bold and "加粗文字" in run.text:
                    has_bold = True
        assert has_bold

    def test_italic_text(self):
        html = "<p><em>斜体文字</em></p>"
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        has_italic = False
        for p in doc.paragraphs:
            for run in p.runs:
                if run.italic and "斜体文字" in run.text:
                    has_italic = True
        assert has_italic

    def test_unordered_list(self):
        html = "<ul><li>项目一</li><li>项目二</li><li>项目三</li></ul>"
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        list_items = [p.text for p in doc.paragraphs if p.style.name == "List Bullet"]
        assert len(list_items) == 3
        assert "项目一" in list_items[0]
        assert "项目二" in list_items[1]
        assert "项目三" in list_items[2]

    def test_ordered_list(self):
        html = "<ol><li>第一步</li><li>第二步</li></ol>"
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        list_items = [p.text for p in doc.paragraphs if p.style.name == "List Number"]
        assert len(list_items) == 2
        assert "第一步" in list_items[0]
        assert "第二步" in list_items[1]

    def test_table(self):
        html = """<table><tr><th>姓名</th><th>年龄</th></tr>
                  <tr><td>张明</td><td>22</td></tr></table>"""
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 2
        assert len(table.columns) == 2
        assert "姓名" in table.rows[0].cells[0].text
        assert "张明" in table.rows[1].cells[0].text

    def test_horizontal_rule(self):
        html = "<p>上方内容</p><hr><p>下方内容</p>"
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "上方内容" in texts
        assert "下方内容" in texts

    def test_div_container(self):
        html = "<div><p>段落一</p><p>段落二</p></div>"
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "段落一" in texts
        assert "段落二" in texts

    def test_nested_div(self):
        html = "<div><div><p>嵌套内容</p></div></div>"
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "嵌套内容" in texts

    def test_link(self):
        html = '<p><a href="https://github.com/test">GitHub</a></p>'
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "GitHub" in texts[0]

    def test_hidden_element(self):
        html = '<div style="display:none;">隐藏内容</div><p>可见内容</p>'
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "可见内容" in texts
        assert "隐藏内容" not in texts

    def test_page_size_a4(self):
        html = "<p>测试</p>"
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        section = doc.sections[0]
        assert abs(section.page_width.cm - 21.0) < 0.1
        assert abs(section.page_height.cm - 29.7) < 0.1
        assert abs(section.top_margin.cm - 1.5) < 0.1
        assert abs(section.left_margin.cm - 1.5) < 0.1

    def test_default_font(self):
        html = "<p>测试字体</p>"
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        style = doc.styles["Normal"]
        assert style.font.name == "微软雅黑"
        assert style.font.size == Pt(10.5)

    def test_theme_color_heading(self):
        html = "<h1>标题</h1>"
        buffer = html_to_word(html, "测试", "template_02")
        doc = Document(buffer)
        for p in doc.paragraphs:
            if p.text.strip() == "标题":
                for run in p.runs:
                    if run.text.strip() == "标题":
                        assert run.font.color.rgb is not None
                        assert str(run.font.color.rgb) == "2563EB"
                break

    def test_empty_html(self):
        html = ""
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        assert len(doc.paragraphs) >= 0

    def test_complex_resume_html(self):
        html = """
        <div class="resume">
            <div style="text-align:center;">
                <h1 style="font-size:24px;">张明</h1>
                <p>13812345678 | zhangming@email.com</p>
            </div>
            <h2>求职意向</h2>
            <p>前端开发工程师 | 北京 | 全职</p>
            <h2>教育背景</h2>
            <p><strong>北京理工大学</strong> | 计算机科学与技术 | 本科</p>
            <h2>技能特长</h2>
            <ul>
                <li>React、Vue、TypeScript</li>
                <li>JavaScript、CSS3、Webpack</li>
            </ul>
            <h2>实习经历</h2>
            <p><strong>前端开发实习生 | 字节跳动 | 2023.06-2023.09</strong></p>
            <ul>
                <li>参与公司内部管理系统开发</li>
                <li>优化页面加载性能</li>
            </ul>
            <h2>自我评价</h2>
            <p>具有扎实的计算机基础知识和Web开发经验。</p>
        </div>
        """
        buffer = html_to_word(html, "张明", "template_15")
        doc = Document(buffer)
        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "张明" in all_text
        assert "求职意向" in all_text
        assert "前端开发工程师" in all_text
        assert "北京理工大学" in all_text
        assert "React" in all_text
        assert "字节跳动" in all_text
        assert "自我评价" in all_text

    def test_no_duplicate_text_in_strong(self):
        html = "<p><strong>加粗内容</strong></p>"
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        for p in doc.paragraphs:
            if "加粗内容" in p.text:
                assert p.text.count("加粗内容") == 1

    def test_inline_font_size(self):
        html = '<p><span style="font-size:18px;">大号文字</span></p>'
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        found = False
        for p in doc.paragraphs:
            for run in p.runs:
                if "大号文字" in run.text and run.font.size:
                    found = True
        assert found

    def test_inline_color(self):
        html = '<p><span style="color:#FF0000;">红色文字</span></p>'
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        found = False
        for p in doc.paragraphs:
            for run in p.runs:
                if "红色文字" in run.text and run.font.color.rgb:
                    found = True
        assert found

    def test_center_alignment(self):
        html = '<p style="text-align:center;">居中文字</p>'
        buffer = html_to_word(html, "测试", "template_15")
        doc = Document(buffer)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        for p in doc.paragraphs:
            if "居中文字" in p.text:
                assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
                break


class TestExportWordAPI:
    def test_export_word_success(self):
        response = client.post("/api/export/word", json={
            "html": "<h1>测试简历</h1><p>内容</p>",
            "name": "测试",
            "templateId": "template_15"
        })
        assert response.status_code == 200
        assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert len(response.content) > 0
        doc = Document(io.BytesIO(response.content))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "测试简历" in texts

    def test_export_word_empty_html(self):
        response = client.post("/api/export/word", json={
            "html": "",
            "name": "测试"
        })
        assert response.status_code == 400
        assert "不能为空" in response.json()["detail"]

    def test_export_word_whitespace_html(self):
        response = client.post("/api/export/word", json={
            "html": "   ",
            "name": "测试"
        })
        assert response.status_code == 400

    def test_export_word_with_chinese_name(self):
        response = client.post("/api/export/word", json={
            "html": "<p>内容</p>",
            "name": "张明",
            "templateId": "template_15"
        })
        assert response.status_code == 200
        assert "content-disposition" in response.headers
        assert "张明" in response.headers["content-disposition"] or "%E5%BC%A0%E6%98%8E" in response.headers["content-disposition"]

    def test_export_word_special_chars_in_name(self):
        response = client.post("/api/export/word", json={
            "html": "<p>内容</p>",
            "name": '测试/<>|"名字',
            "templateId": "template_15"
        })
        assert response.status_code == 200

    def test_export_word_default_template(self):
        response = client.post("/api/export/word", json={
            "html": "<p>内容</p>",
            "name": "测试"
        })
        assert response.status_code == 200

    def test_export_word_different_templates(self):
        for tid in ["template_01", "template_06", "template_15"]:
            response = client.post("/api/export/word", json={
                "html": "<h1>标题</h1>",
                "name": "测试",
                "templateId": tid
            })
            assert response.status_code == 200

    def test_export_word_complex_html(self):
        html = """
        <div>
            <h1>张明</h1>
            <h2>教育背景</h2>
            <p><strong>北京理工大学</strong> | 计算机 | 本科</p>
            <h2>技能</h2>
            <ul><li>React</li><li>Vue</li></ul>
            <h2>经历</h2>
            <table><tr><th>公司</th><th>职位</th></tr>
            <tr><td>字节</td><td>前端</td></tr></table>
        </div>
        """
        response = client.post("/api/export/word", json={
            "html": html,
            "name": "张明",
            "templateId": "template_02"
        })
        assert response.status_code == 200
        doc = Document(io.BytesIO(response.content))
        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "张明" in all_text
        assert "教育背景" in all_text
        assert "北京理工大学" in all_text

    def test_export_word_valid_docx_structure(self):
        response = client.post("/api/export/word", json={
            "html": "<p>验证docx结构</p>",
            "name": "测试",
            "templateId": "template_15"
        })
        assert response.status_code == 200
        buffer = io.BytesIO(response.content)
        doc = Document(buffer)
        assert len(doc.sections) > 0
        section = doc.sections[0]
        assert abs(section.page_width.cm - 21.0) < 0.1
