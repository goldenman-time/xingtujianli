from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, validator
from typing import List, Optional
import os
import uuid
import time
import hashlib
from datetime import datetime
from dotenv import load_dotenv
import json
import httpx
import asyncio
import re
import shutil
import io
import tempfile
import urllib.parse
from html import escape as html_escape

load_dotenv()

# ========== 后端XSS安全防护工具 ==========

class SecurityUtils:
    """统一的安全工具类，用于防止XSS和其他注入攻击"""
    
    @staticmethod
    def escape_html(text: str) -> str:
        """转义HTML特殊字符"""
        if not isinstance(text, str):
            return str(text) if text else ''
        return html_escape(text)
    
    @staticmethod
    def sanitize_input(text: str, max_length: int = 10000) -> str:
        """净化用户输入，移除潜在的危险内容"""
        if not isinstance(text, str):
            return str(text) if text else ''
        
        # 限制长度
        text = text[:max_length]
        
        # 移除null字节
        text = text.replace('\x00', '')
        
        return text
    
    @staticmethod
    def validate_name(name: str) -> tuple:
        """验证姓名输入"""
        if not name or not name.strip():
            return False, name, "姓名不能为空"
        
        name = name.strip()
        if len(name) > 50:
            return False, name[:50], "姓名过长"
        
        # 只允许中文、字母、空格和常见分隔符
        if not re.match(r'^[\u4e00-\u9fa5a-zA-Z\s·\-\.]+$', name):
            # 移除非法字符
            clean_name = re.sub(r'[^\u4e00-\u9fa5a-zA-Z\s·\-\.]', '', name)
            return False, clean_name, "姓名包含非法字符"
        
        return True, name, ""
    
    @staticmethod
    def validate_phone(phone: str) -> tuple:
        """验证手机号"""
        if not phone:
            return False, "", "手机号不能为空"
        
        phone = phone.strip()
        clean_phone = re.sub(r'[\s\-—\(\)]', '', phone)
        
        if not re.match(r'^1[3-9]\d{9}$', clean_phone):
            return False, phone, "手机号格式不正确"
        
        return True, clean_phone, ""
    
    @staticmethod
    def validate_email(email: str) -> tuple:
        """验证邮箱"""
        if not email:
            return False, "", "邮箱不能为空"
        
        email = email.strip()
        if not re.match(r'^[^\s@]+@[^\s@]+\.[^\s@]+$', email):
            return False, email, "邮箱格式不正确"
        
        if len(email) > 100:
            return False, email[:100], "邮箱过长"
        
        return True, email, ""
    
    @staticmethod
    def validate_url(url: str) -> tuple:
        """验证URL"""
        if not url or not url.strip():
            return True, "", ""  # URL是可选的
        
        url = url.strip()
        try:
            parsed = urllib.parse.urlparse(url)
            if parsed.scheme not in ['http', 'https']:
                return False, "", "URL必须以http或https开头"
            
            if len(url) > 500:
                return False, url[:500], "URL过长"
            
            return True, url, ""
        except Exception:
            return False, "", "URL格式不正确"
    
    @staticmethod
    def validate_text_content(content: str, field_name: str = "内容", max_length: int = 10000) -> tuple:
        """验证文本内容"""
        if not content:
            return True, "", ""
        
        content = content.strip()
        
        if len(content) > max_length:
            return False, content[:max_length], f"{field_name}过长（最多{max_length}字）"
        
        # 检查是否包含可疑的脚本标签（虽然会在输出时转义，但提前检测更好）
        dangerous_patterns = [
            r'<script[^>]*>.*?</script>',
            r'javascript:',
            r'on\w+\s*=',
            r'data:text/html',
        ]
        
        for pattern in dangerous_patterns:
            if re.search(pattern, content, re.IGNORECASE | re.DOTALL):
                # 记录警告但不拒绝（因为会转义）
                pass
        
        return True, content, ""
    
    @staticmethod
    def sanitize_generated_html(html: str) -> str:
        """净化AI生成的HTML内容，移除潜在的危险标签和属性"""
        if not isinstance(html, str):
            return ''
        
        # 移除script标签及其内容
        html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.IGNORECASE | re.DOTALL)
        
        # 移除事件处理器
        html = re.sub(r'\s+on\w+\s*=\s*["\'][^"\']*["\']', '', html, flags=re.IGNORECASE)
        
        # 移除javascript:协议
        html = re.sub(r'javascript\s*:', '', html, flags=re.IGNORECASE)
        
        # 移除data: URL（除了图片）
        html = re.sub(r'data:(?!image/)[^"\'>\s]*', '', html, flags=re.IGNORECASE)
        
        # 移除危险标签（保留安全的格式化标签）
        dangerous_tags = ['script', 'iframe', 'object', 'embed', 'form', 'input', 'button']
        for tag in dangerous_tags:
            html = re.sub(f'<{tag}[^>]*>.*?</{tag}>', '', html, flags=re.IGNORECASE | re.DOTALL)
            html = re.sub(f'<{tag}[^>]*/?>', '', html, flags=re.IGNORECASE)
        
        return html.strip()

app = FastAPI(title="星途简历 - AI后端", version="3.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

API_KEY = os.getenv("DOUBAO_API_KEY")
MODEL_ID = os.getenv("DOUBAO_MODEL_ID", "doubao-pro-32k-241215")
BASE_URL = "https://ark.cn-beijing.volces.com/api/v3/chat/completions"

# ========== 数据模型 ==========

class ExperienceItem(BaseModel):
    id: int
    type: str
    role: str
    company: str
    date: str
    content: str

class StarResult(BaseModel):
    experienceId: int
    experienceType: str
    role: str
    company: str
    date: str
    focus: str
    S: str
    T: str
    A: str
    R: str

class SuggestionItem(BaseModel):
    type: str
    icon: str
    label: str
    subject: str
    issue: str
    direction: str
    example: str

class SuggestionsResponse(BaseModel):
    title: str
    isExcellent: bool
    suggestions: List[SuggestionItem]

# ========== 新增：完整字段对齐前端 ==========

class OptimizeRequest(BaseModel):
    targetJob: str = ""
    expectedCity: str = ""
    jobType: str = "fulltime"
    name: str = ""
    phone: str = ""
    email: str = ""
    photo: str = ""
    github: str = ""
    blog: str = ""
    portfolio: str = ""
    personalWebsite: str = ""
    school: str = ""
    major: str = ""
    educationLevel: str = "undergraduate"
    graduation: str = ""
    gpa: str = ""
    courses: str = ""
    skills: str = ""
    certificates: str = ""
    selfEvaluation: str = ""
    jd: str = ""
    templateId: str = "template_15"
    experiences: List[ExperienceItem]

class OptimizeResponse(BaseModel):
    results: List[StarResult]

class SuggestionsRequest(BaseModel):
    experience: ExperienceItem
    jd: str = ""
    targetJob: str = ""
    skills: str = ""

class GenerateResumeRequest(BaseModel):
    targetJob: str = ""
    expectedCity: str = ""
    jobType: str = "fulltime"
    name: str = ""
    phone: str = ""
    email: str = ""
    photo: str = ""
    github: str = ""
    blog: str = ""
    portfolio: str = ""
    personalWebsite: str = ""
    school: str = ""
    major: str = ""
    educationLevel: str = "undergraduate"
    graduation: str = ""
    gpa: str = ""
    courses: str = ""
    skills: str = ""
    certificates: str = ""
    selfEvaluation: str = ""
    jd: str = ""
    templateId: str = "template_15"
    experiences: List[ExperienceItem]
    starResults: List[StarResult]
    suggestions: Optional[List[dict]] = None

class GenerateResumeResponse(BaseModel):
    html: str
    success: bool

class ExportWordRequest(BaseModel):
    html: str
    name: str = "简历"
    templateId: str = "template_15"

# ========== HTML转Word文档 ==========

def html_to_word(html_content: str, name: str = "简历", template_id: str = "template_15") -> io.BytesIO:
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn

    tmpl = TEMPLATE_STYLES.get(template_id, TEMPLATE_STYLES["template_15"])
    theme_color_hex = tmpl.get("color", "#3b82f6").lstrip("#")
    theme_color = RGBColor(
        int(theme_color_hex[0:2], 16),
        int(theme_color_hex[2:4], 16),
        int(theme_color_hex[4:6], 16)
    )

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    pf = style.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.35

    soup = BeautifulSoup(html_content, "lxml")

    def _parse_style(style_str):
        result = {}
        if not style_str:
            return result
        for item in style_str.split(";"):
            item = item.strip()
            if ":" in item:
                k, v = item.split(":", 1)
                result[k.strip().lower()] = v.strip()
        return result

    def _get_font_size(styles):
        fs = styles.get("font-size", "")
        if not fs:
            return None
        fs = fs.lower().replace(" ", "")
        try:
            if "px" in fs:
                val = float(fs.replace("px", ""))
                return Pt(val * 0.75)
            elif "pt" in fs:
                return Pt(float(fs.replace("pt", "")))
            elif "em" in fs:
                return Pt(float(fs.replace("em", "")) * 10.5)
        except (ValueError, TypeError):
            pass
        return None

    def _get_color(styles):
        c = styles.get("color", "")
        if not c:
            return None
        c = c.strip().lower()
        if c.startswith("#") and len(c) == 7:
            try:
                return RGBColor(int(c[1:3], 16), int(c[3:5], 16), int(c[5:7], 16))
            except ValueError:
                pass
        return None

    def _add_inline_runs(paragraph, element, default_bold=False):
        if element.name == "br":
            paragraph.add_run("\n")
            return
        if isinstance(element, str):
            text = element.strip()
            if text:
                run = paragraph.add_run(text)
                run.bold = default_bold
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            return
        if not hasattr(element, 'name') or element.name is None:
            text = str(element).strip()
            if text:
                run = paragraph.add_run(text)
                run.bold = default_bold
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            return

        tag = element.name.lower() if element.name else ""
        el_style = _parse_style(element.get("style", ""))
        is_bold = default_bold or tag in ("b", "strong")
        is_italic = tag in ("i", "em")
        is_underline = tag in ("u",)
        is_strikethrough = tag in ("s", "strike", "del")
        font_weight = el_style.get("font-weight", "")
        if font_weight in ("bold", "700", "800", "900"):
            is_bold = True
        font_style_css = el_style.get("font-style", "")
        if font_style_css == "italic":
            is_italic = True
        text_decoration = el_style.get("text-decoration", "")
        if "underline" in text_decoration:
            is_underline = True
        if "line-through" in text_decoration:
            is_strikethrough = True

        if tag in ("img", "br", "hr"):
            if tag == "br":
                paragraph.add_run("\n")
            return

        if tag in ("ul", "ol"):
            return

        has_element_children = any(hasattr(c, 'name') and c.name is not None for c in element.children)

        if has_element_children:
            for child in element.children:
                _add_inline_runs(paragraph, child, is_bold)
        else:
            text = element.get_text()
            if text and text.strip():
                run = paragraph.add_run(text.strip())
                run.bold = is_bold
                run.italic = is_italic
                run.underline = is_underline
                if is_strikethrough:
                    run.font.strike = True
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                font_size = _get_font_size(el_style)
                if font_size:
                    run.font.size = font_size
                color = _get_color(el_style)
                if color:
                    run.font.color.rgb = color

    def _process_element(element, doc):
        if isinstance(element, str):
            text = element.strip()
            if text:
                p = doc.add_paragraph(text)
            return

        if not hasattr(element, 'name') or element.name is None:
            return

        tag = element.name.lower() if element.name else ""

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            text = element.get_text(strip=True)
            if not text:
                return
            p = doc.add_paragraph()
            if level <= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.bold = True
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            size_map = {1: Pt(22), 2: Pt(16), 3: Pt(14), 4: Pt(12), 5: Pt(11), 6: Pt(10.5)}
            run.font.size = size_map.get(level, Pt(10.5))
            if level <= 2:
                run.font.color.rgb = theme_color
            pf = p.paragraph_format
            pf.space_before = Pt(12) if level <= 2 else Pt(6)
            pf.space_after = Pt(4)
            if level <= 2:
                p2 = doc.add_paragraph()
                p2_fmt = p2.paragraph_format
                p2_fmt.space_before = Pt(0)
                p2_fmt.space_after = Pt(6)
                from docx.oxml import OxmlElement
                pPr = p2._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), theme_color_hex)
                pBdr.append(bottom)
                pPr.append(pBdr)
            return

        if tag == "p":
            p = doc.add_paragraph()
            el_style = _parse_style(element.get("style", ""))
            align = el_style.get("text-align", "")
            if align == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for child in element.children:
                _add_inline_runs(p, child)
            return

        if tag == "ul":
            for li in element.find_all("li", recursive=False):
                text = li.get_text(strip=True)
                if text:
                    p = doc.add_paragraph(style='List Bullet')
                    for child in li.children:
                        _add_inline_runs(p, child, False)
            return

        if tag == "ol":
            for li in element.find_all("li", recursive=False):
                text = li.get_text(strip=True)
                if text:
                    p = doc.add_paragraph(style='List Number')
                    for child in li.children:
                        _add_inline_runs(p, child, False)
            return

        if tag == "table":
            rows = element.find_all("tr")
            if not rows:
                return
            first_row = rows[0]
            is_header = bool(first_row.find(["th"]))
            num_cols = len(first_row.find_all(["td", "th"]))
            if num_cols == 0:
                return
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'
            for i, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for j, cell in enumerate(cells):
                    if j < num_cols:
                        table_cell = table.rows[i].cells[j]
                        cell_text = cell.get_text(strip=True)
                        table_cell.text = ""
                        p = table_cell.paragraphs[0]
                        for child in cell.children:
                            _add_inline_runs(p, child)
                        if i == 0 and is_header:
                            for run in p.runs:
                                run.bold = True
            return

        if tag == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            from docx.oxml import OxmlElement
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            return

        if tag in ("div", "section", "article", "main", "header", "footer", "nav", "aside", "span"):
            block_style = _parse_style(element.get("style", ""))
            display = block_style.get("display", "")
            if display == "none":
                return
            for child in element.children:
                _process_element(child, doc)
            return

        if tag in ("a",):
            text = element.get_text(strip=True)
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
                run.underline = True
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                href = element.get("href", "")
                if href:
                    run2 = p.add_run(f" ({href})")
                    run2.font.size = Pt(8)
                    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                    run2.font.name = '微软雅黑'
                    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            return

        for child in element.children:
            _process_element(child, doc)

    body = soup.body or soup
    for child in body.children:
        _process_element(child, doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ========== 模板风格映射 ==========

TEMPLATE_STYLES = {
    "template_01": {"name": "经典商务单栏", "style_desc": "简约商务风，ATS兼容性100%，适合财务/HR/行政/管培生", "color": "#1a365d", "layout": "单栏", "font": "微软雅黑/Arial"},
    "template_02": {"name": "互联网技术双栏", "style_desc": "互联网风，技术岗专属，技能前置展示，左右分栏(3:7)，适合Java/前端/后端/算法", "color": "#2563eb", "layout": "左右分栏(3:7)", "font": "Inter/思源黑体"},
    "template_03": {"name": "金融投行精英", "style_desc": "简约商务风，极致专业，信息密度极高，单栏紧凑，适合投行/研究员/PE/VC/四大", "color": "#000000", "layout": "单栏紧凑", "font": "Times New Roman/宋体"},
    "template_04": {"name": "产品经理思维", "style_desc": "互联网风，突出产品思维和数据成果，适合产品经理/产品运营/B端产品", "color": "#f97316", "layout": "单栏", "font": "现代无衬线字体"},
    "template_05": {"name": "运营增长数据", "style_desc": "互联网风，数据驱动，增长指标突出，单栏+数据卡片，适合运营/增长/新媒体/营销", "color": "#10b981", "layout": "单栏+数据卡片", "font": "清晰无衬线字体"},
    "template_06": {"name": "国企央企正式", "style_desc": "国企/体制内风，政治面貌前置，稳重正式，表格化单栏，适合公务员/事业单位/国企/央企", "color": "#dc2626", "layout": "表格化单栏", "font": "宋体/仿宋"},
    "template_07": {"name": "创意设计作品", "style_desc": "创意设计风，设计感第一，作品展示充分，左侧信息栏+右侧作品区，适合UI/UX/平面/视觉设计", "color": "#8b5cf6", "layout": "左侧信息栏+右侧作品区", "font": "多种字体组合"},
    "template_08": {"name": "学术科研专业", "style_desc": "学术科研风，LaTeX风格，论文为核心，学术论文式单栏，适合博士申请/高校教师/研究员", "color": "#000000", "layout": "学术论文式单栏", "font": "Times New Roman/Computer Modern"},
    "template_09": {"name": "市场营销品牌", "style_desc": "创意商务风，Campaign案例突出，品牌思维，现代商务单栏，适合品牌经理/市场推广/公关/广告", "color": "#ef4444", "layout": "现代商务单栏", "font": "专业商务字体"},
    "template_10": {"name": "医疗健康专业", "style_desc": "专业严谨风，资质证书前置，临床经验，规整单栏，适合医生/护士/医药代表/医疗行政", "color": "#059669", "layout": "规整单栏", "font": "清晰标准字体"},
    "template_11": {"name": "制造业工程", "style_desc": "专业技术风，技术认证突出，产线经验，结构化单栏，适合工艺/机械工程师/生产管理/质控", "color": "#1e40af", "layout": "结构化单栏", "font": "标准技术文档字体"},
    "template_12": {"name": "教育培训行业", "style_desc": "专业亲和风，教学成果突出，亲和力强，整洁单栏，适合教师/教研/课程设计/培训师", "color": "#3b82f6", "layout": "整洁单栏", "font": "清晰易读字体"},
    "template_13": {"name": "行政人事管理", "style_desc": "稳重商务风，协调能力突出，流程优化，清晰规整单栏，适合行政/HR/办公室主任", "color": "#4b5563", "layout": "清晰规整单栏", "font": "标准商务字体"},
    "template_14": {"name": "咨询管理精英", "style_desc": "高端商务风，麦肯锡风格，极度结构化，高密度单栏，适合管理咨询/战略咨询/IT咨询", "color": "#1e3a8a", "layout": "高密度单栏", "font": "Times New Roman/宋体"},
    "template_15": {"name": "应届生校招全能", "style_desc": "简约商务风，应届生专属，ATS友好，标准单栏，适合应届生校招/通用/无明确方向", "color": "#3b82f6", "layout": "标准单栏", "font": "微软雅黑"},
    "template_16": {"name": "社招3-5年进阶", "style_desc": "专业商务风，职业成长清晰，能力层级明确，经历排序单栏，适合3-5年经验/主管/高级专员", "color": "#0369a1", "layout": "经历排序单栏", "font": "成熟稳重字体"},
    "template_17": {"name": "社招5年+管理", "style_desc": "高端商务风，管理能力清晰，战略高度，大气双栏/单栏，适合5年+经验/经理/总监", "color": "#1e40af", "layout": "大气双栏/单栏", "font": "权威感字体"},
    "template_18": {"name": "转行跨界桥梁", "style_desc": "能力迁移风，可迁移能力前置，衔接新旧职业，能力模块前置单栏，适合跨行业求职/职业转型", "color": "#7c3aed", "layout": "能力模块前置单栏", "font": "专业清晰字体"},
}

# ========== 调用豆包 AI ==========

async def call_doubao_optimize(exp: ExperienceItem, jd: str, school: str, major: str) -> StarResult:
    system_prompt = """你是一位拥有10年经验的资深HR总监和简历优化专家。
你的任务是根据用户提供的原始经历描述，使用 STAR 法则进行专业优化。

【STAR 法则说明】
- S (Situation/情境): 描述当时的背景、环境、面临的挑战
- T (Task/任务): 明确你的具体职责、目标、需要解决的问题
- A (Action/行动): 详细说明你采取的具体措施、方法、使用的技能
- R (Result/结果): 用数据量化最终成果，如提升了XX%、完成了XX、获得了XX

【重要规则】
1. 必须基于用户提供的真实经历，严禁虚构数据或编造成果
2. 每个要素控制在 50-80 字，语言简洁专业
3. 结果部分必须有具体数字
4. 输出必须是严格的 JSON 格式，不要包含 markdown 代码块标记

【输出格式】
{
    "S": "情境描述...",
    "T": "任务描述...",
    "A": "行动描述...",
    "R": "结果描述（必须包含数字）..."
}"""

    user_prompt = f"""【经历类型】{exp.type}
【职位/角色】{exp.role}
【公司/机构】{exp.company}
【时间段】{exp.date}
【原始经历描述】{exp.content}
【学校背景】{school} {major}
【目标岗位JD】{jd if jd else "未提供"}

请按 STAR 法则优化这段经历，只输出 JSON 格式，不要添加任何其他文字。"""

    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.post(
            BASE_URL,
            headers={"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"},
            json={
                "model": MODEL_ID,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 800
            }
        )
        result = response.json()
        ai_content = clean_json_response(result["choices"][0]["message"]["content"])
        star_data = json.loads(ai_content)

        focus_map = {
            "internship": "突出工作成果与业务能力",
            "campus": "突出技术能力与问题解决",
            "club": "突出组织协调与领导力",
            "parttime": "突出责任心与执行力",
            "research": "突出科研能力与学术成果"
        }

        return StarResult(
            experienceId=exp.id,
            experienceType=exp.type,
            role=exp.role,
            company=exp.company,
            date=exp.date,
            focus=focus_map.get(exp.type, "综合优化"),
            S=star_data.get("S", ""),
            T=star_data.get("T", ""),
            A=star_data.get("A", ""),
            R=star_data.get("R", "")
        )

async def call_doubao_suggestions(exp: ExperienceItem, jd: str, target_job: str = "", skills: str = "") -> SuggestionsResponse:
    system_prompt = """你是一位拥有10年经验的资深HR总监和简历优化专家。
你的任务是根据用户提供的原始经历、经历类型和目标岗位JD，生成个性化、可落地的优化建议。

【分析流程】（必须在后台完成，不输出分析过程）
1. STAR法则完整性分析：检查S/T/A/R四要素覆盖情况
2. 岗位需求匹配度分析：从JD提取关键词，标记经历中未提及的技能点
3. 成果量化程度分析：识别"提升效率"、"效果很好"、"参与负责"等模糊表述
4. 经历类型适配性分析：根据类型（实习/校园项目/社团/兼职）检查特性描述
5. 个人贡献清晰度分析：检查"参与"、"协助"等弱化个人贡献的表述

【输出格式要求】
必须输出严格的JSON格式，不要包含任何markdown标记：

{
    "isExcellent": false,
    "suggestions": [
        {
            "type": "critical",
            "icon": "⚠️",
            "label": "核心问题",
            "subject": "建议标题（15字以内）",
            "issue": "【问题定位】具体指出原始经历中的问题，必须引用用户原文中的具体表述",
            "direction": "【具体修改方向】提供可操作的优化方法",
            "example": "【参考示例】基于用户实际经历优化后的表述"
        }
    ]
}

【规则】
1. 所有建议必须直接引用用户原始经历中的具体表述，禁止通用模板
2. 建议总数1-5条，按严重性排序：核心问题 > 岗位适配 > 加分项
3. 如果经历非常优秀（STAR完整、匹配度≥85%、量化≥70%、贡献清晰），isExcellent设为true，只输出2条加分项
4. 参考示例必须基于用户实际经历，不得捏造不存在的内容
5. 禁止使用"建议优化内容结构"、"建议突出核心能力"等通用表述
6. 问题定位必须精确到用户原文中的具体词语或句子"""

    user_prompt = f"""【经历类型】{exp.type}
【职位/角色】{exp.role}
【公司/机构】{exp.company}
【原始经历文本】{exp.content}
【目标岗位】{target_job if target_job else "未提供"}
【目标岗位JD】{jd if jd else "未提供"}
【已掌握技能】{skills if skills else "未提供"}

请生成个性化优化建议，只输出JSON格式。"""

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                BASE_URL,
                headers={"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"},
                json={
                    "model": MODEL_ID,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    "temperature": 0.7,
                    "max_tokens": 1500
                }
            )
            response.raise_for_status()
            result = response.json()
            ai_content = clean_json_response(result["choices"][0]["message"]["content"])
            data = json.loads(ai_content)

            suggestions = []
            for item in data.get("suggestions", []):
                suggestions.append(SuggestionItem(
                    type=item.get("type", "bonus"),
                    icon=item.get("icon", "💡"),
                    label=item.get("label", "加分项"),
                    subject=item.get("subject", ""),
                    issue=item.get("issue", ""),
                    direction=item.get("direction", ""),
                    example=item.get("example", "")
                ))

            is_excellent = data.get("isExcellent", False)
            count = len(suggestions)
            title = f"🎉 你的经历描述非常优秀！以下是2个可以锦上添花的小建议" if is_excellent else f"📝 你的简历个性化优化建议（共{count}条）"

            return SuggestionsResponse(title=title, isExcellent=is_excellent, suggestions=suggestions)
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=502, detail=f"AI服务请求失败({e.response.status_code})")
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="AI服务响应超时，请稍后重试")
    except json.JSONDecodeError:
        return SuggestionsResponse(
            title="⚠️ AI返回数据解析异常",
            isExcellent=False,
            suggestions=[SuggestionItem(
                type="bonus", icon="💡", label="通用建议",
                subject="AI解析异常", issue="", direction="建议稍后重试获取更精准的建议", example=""
            )]
        )
    except Exception as e:
        import logging
        logging.error(f"Suggestions API error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"生成建议时出错: {str(e)}")

async def call_doubao_generate_resume(req: GenerateResumeRequest) -> str:
    tmpl = TEMPLATE_STYLES.get(req.templateId, TEMPLATE_STYLES["template_15"])
    
    star_text = ""
    for r in req.starResults:
        star_text += f"\n【{r.role} | {r.company} | {r.date}】\nS: {r.S}\nT: {r.T}\nA: {r.A}\nR: {r.R}\n"

    exp_text = ""
    for e in req.experiences:
        exp_text += f"\n【{e.role} | {e.company}】\n{e.content}\n"

    suggestion_text = ""
    if req.suggestions:
        for idx, group in enumerate(req.suggestions, 1):
            suggestion_text += f"\n--- 建议组{idx} ---\n"
            for s in group.get("suggestions", []):
                suggestion_text += f"- [{s.get('label','')}] {s.get('subject','')}: {s.get('direction','')}\n"

    system_prompt = f"""你是一位拥有15年经验的顶级HR总监、简历优化专家和前端设计师。
你的任务是根据用户提供的所有信息，生成一份完整、专业、可直接投递使用的简历HTML。

【当前选中的简历模板风格】
模板名称：{tmpl['name']}
风格描述：{tmpl['style_desc']}
主色调：{tmpl['color']}
布局：{tmpl['layout']}
推荐字体：{tmpl['font']}

【设计规范】
- 必须严格遵循上述模板风格，配色、布局、字体气质要与模板描述一致
- 如果是技术岗模板（如互联网双栏），技能板块要前置，使用左右分栏布局
- 如果是金融/咨询模板，信息密度要高，使用Times New Roman类字体，排版紧凑
- 如果是创意/设计模板，可以增加设计感元素，但不要过度花哨
- 如果是国企/体制内模板，风格稳重正式，使用宋体/仿宋
- 通用要求：适配A4打印，内容宽度最大210mm，内边距15mm，确保可打印在1-2页A4纸内

【内容规范】
1. 个人信息：姓名用超大字号(24px+)突出，联系方式横排，GitHub/Blog/Portfolio/个人网站做成可点击链接（如有）
2. 求职意向：目标岗位、期望城市、工作类型要清晰展示
3. 教育背景：学校、专业、时间、学历、GPA（如有）、主修课程（精简为1行）
4. 技能特长：必须展示用户填写的技能和证书，分点列出，使用熟练度描述（熟悉/掌握/了解）
5. 经历部分：必须使用优化后的STAR内容，突出数据量化成果，每条经历整合成连贯的3-4句话，不要分S/T/A/R小标题
6. 自我评价：3-4句话，针对目标岗位JD定制，如用户未填写则根据经历自动生成
7. 所有内容必须真实，基于用户提供的信息，严禁虚构任何经历或数据

【HTML规范】
1. 输出完整HTML文档片段（从<body>内的内容开始即可，不需要html/head标签，因为会嵌入到现有页面）
2. CSS使用内联style属性，确保样式不丢失
3. 使用中文，排版整齐，行高1.6，段间距合理
4. 不要包含任何JavaScript
5. 只输出HTML代码，不要任何解释、不要markdown代码块标记
6. 确保所有颜色对比度足够，打印清晰"""

    user_prompt = f"""请生成完整简历HTML内容。

【基本信息】
姓名：{req.name}
电话：{req.phone}
邮箱：{req.email}
照片：{req.photo or '未提供'}
GitHub：{req.github or '未提供'}
博客：{req.blog or '未提供'}
作品集：{req.portfolio or '未提供'}
个人网站：{req.personalWebsite or '未提供'}

【求职意向】
目标岗位：{req.targetJob or '未提供'}
期望城市：{req.expectedCity or '未提供'}
工作类型：{req.jobType or '未提供'}

【教育背景】
学校：{req.school}
专业：{req.major}
学历：{req.educationLevel}
毕业时间：{req.graduation}
GPA：{req.gpa or '未提供'}
主修课程：{req.courses or '未提供'}

【技能特长】
专业技能：{req.skills or '未提供'}
证书/奖项：{req.certificates or '未提供'}

【自我评价】
{req.selfEvaluation or '未提供，请根据经历自动生成'}

【目标岗位JD】
{req.jd or '未提供'}

【优化后的经历（STAR法则）】
{star_text}

【原始经历】
{exp_text}

【优化建议】（请确保这些建议已被采纳到简历中）
{suggestion_text}

请直接输出简历的HTML body内容（不需要<html><head>，只需要<div class="resume">...</div>这样的内容块）。"""

    async with httpx.AsyncClient(timeout=120.0) as client:
        try:
            response = await client.post(
                BASE_URL,
                headers={"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"},
                json={
                    "model": MODEL_ID,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    "temperature": 0.6,
                    "max_tokens": 3000
                }
            )
            response.raise_for_status()
            result = response.json()
            ai_content = result["choices"][0]["message"]["content"]
            html = clean_json_response(ai_content)
        except httpx.HTTPStatusError as e:
            raise HTTPException(status_code=502, detail=f"AI服务请求失败({e.response.status_code})")
        except httpx.TimeoutException:
            raise HTTPException(status_code=504, detail="AI生成简历超时，请稍后重试")
        except (KeyError, IndexError) as e:
            import logging
            logging.error(f"AI返回数据格式异常: {str(e)}", exc_info=True)
            raise HTTPException(status_code=502, detail="AI服务返回数据格式异常")

    if "<body>" in html:
        match = re.search(r'<body[^>]*>(.*?)</body>', html, re.DOTALL)
        if match:
            html = match.group(1)
    return html.strip()

def clean_json_response(content: str) -> str:
    content = content.strip()
    if content.startswith("```json"):
        content = content[7:]
    elif content.startswith("```html"):
        content = content[8:]
    elif content.startswith("```"):
        content = content[3:]
    if content.endswith("```"):
        content = content[:-3]
    return content.strip()

# ========== 简历上传模块 ==========

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt"}
MAX_FILE_SIZE = 15 * 1024 * 1024

ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}

resume_store = {}

def get_file_extension(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()

def validate_file_type(filename: str, content_type: str) -> bool:
    ext = get_file_extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        return False
    if content_type and content_type not in ALLOWED_MIME_TYPES and not content_type.startswith("application/"):
        if content_type not in ALLOWED_MIME_TYPES:
            return False
    return True

def extract_text_from_file(file_path: str, extension: str) -> str:
    try:
        if extension == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        elif extension == ".pdf":
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
        elif extension == ".docx":
            from docx import Document
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            return text.strip()
        elif extension == ".doc":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            text = re.sub(r'[^\u4e00-\u9fff\u0020-\u007ea-zA-Z0-9\s\.\,\;\:\!\?\-\(\)\[\]\{\}@\/\\]', '', content)
            return text.strip() if text.strip() else "DOC格式文件，建议转换为DOCX格式以获得更好的解析效果"
        else:
            return ""
    except Exception as e:
        return f"文件解析失败: {str(e)}"

def generate_safe_filename(original_filename: str) -> str:
    ext = get_file_extension(original_filename)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = uuid.uuid4().hex[:8]
    return f"resume_{timestamp}_{unique_id}{ext}"

class ResumeChatRequest(BaseModel):
    resume_id: str
    message: str
    chat_history: Optional[List[dict]] = None

# ========== API 路由 ==========

@app.post("/api/optimize", response_model=OptimizeResponse)
async def optimize_resume(req: OptimizeRequest):
    if not API_KEY:
        raise HTTPException(status_code=500, detail="API Key 未配置")
    if not req.experiences:
        raise HTTPException(status_code=400, detail="经历列表不能为空")

    # 验证和净化用户输入
    validated_experiences = []
    for exp in req.experiences:
        # 净化每个经历字段
        safe_exp = ExperienceItem(
            id=exp.id,
            type=exp.type if exp.type in ['internship', 'campus', 'club', 'parttime', 'research'] else 'internship',
            role=SecurityUtils.sanitize_input(exp.role, 100),
            company=SecurityUtils.sanitize_input(exp.company, 100),
            date=SecurityUtils.sanitize_input(exp.date, 50),
            content=SecurityUtils.sanitize_input(exp.content, 5000)
        )
        validated_experiences.append(safe_exp)

    # 净化其他字段
    safe_jd = SecurityUtils.sanitize_input(req.jd, 5000)
    safe_school = SecurityUtils.sanitize_input(req.school, 100)
    safe_major = SecurityUtils.sanitize_input(req.major, 100)

    tasks = [call_doubao_optimize(exp, safe_jd, safe_school, safe_major) for exp in validated_experiences]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    valid_results = []
    for i, result in enumerate(results):
        if isinstance(result, Exception):
            exp = validated_experiences[i]
            valid_results.append(StarResult(
                experienceId=exp.id,
                experienceType=exp.type,
                role=SecurityUtils.escape_html(exp.role),
                company=SecurityUtils.escape_html(exp.company),
                date=SecurityUtils.escape_html(exp.date),
                focus="综合优化（AI处理异常）",
                S=f"在{SecurityUtils.escape_html(exp.company) or '相关单位'}开展实践。",
                T=f"担任{SecurityUtils.escape_html(exp.role)}，负责核心工作任务。",
                A="运用专业技能完成各项工作。",
                R="工作顺利完成，获得认可。"
            ))
        else:
            # 确保AI返回的结果也是安全的
            valid_results.append(StarResult(
                experienceId=result.experienceId,
                experienceType=result.experienceType,
                role=result.role,
                company=result.company,
                date=result.date,
                focus=SecurityUtils.escape_html(result.focus),
                S=SecurityUtils.escape_html(result.S),
                T=SecurityUtils.escape_html(result.T),
                A=SecurityUtils.escape_html(result.A),
                R=SecurityUtils.escape_html(result.R)
            ))

    return OptimizeResponse(results=valid_results)

@app.post("/api/suggestions", response_model=SuggestionsResponse)
async def get_suggestions(req: SuggestionsRequest):
    if not API_KEY:
        raise HTTPException(status_code=500, detail="API Key 未配置")
    if not req.experience.content or req.experience.content.strip() == "":
        raise HTTPException(status_code=400, detail="经历内容不能为空")
    
    # 净化输入
    safe_experience = ExperienceItem(
        id=req.experience.id,
        type=req.experience.type,
        role=SecurityUtils.sanitize_input(req.experience.role, 100),
        company=SecurityUtils.sanitize_input(req.experience.company, 100),
        date=SecurityUtils.sanitize_input(req.experience.date, 50),
        content=SecurityUtils.sanitize_input(req.experience.content, 5000)
    )
    safe_jd = SecurityUtils.sanitize_input(req.jd, 5000)
    safe_target_job = SecurityUtils.sanitize_input(req.targetJob, 100)
    safe_skills = SecurityUtils.sanitize_input(req.skills, 1000)
    
    return await call_doubao_suggestions(safe_experience, safe_jd, safe_target_job, safe_skills)

@app.post("/api/generate-resume", response_model=GenerateResumeResponse)
async def generate_full_resume(req: GenerateResumeRequest):
    if not API_KEY:
        raise HTTPException(status_code=500, detail="API Key 未配置")
    if not req.name or not req.experiences:
        raise HTTPException(status_code=400, detail="姓名和经历不能为空")
    
    try:
        # 净化所有输入字段
        safe_req = GenerateResumeRequest(
            targetJob=SecurityUtils.sanitize_input(req.targetJob, 100),
            expectedCity=SecurityUtils.sanitize_input(req.expectedCity, 50),
            jobType=req.jobType,
            name=SecurityUtils.sanitize_input(req.name, 50),
            phone=SecurityUtils.sanitize_input(req.phone, 20),
            email=SecurityUtils.sanitize_input(req.email, 100),
            photo=SecurityUtils.sanitize_input(req.photo, 500) if req.photo else "",
            github=SecurityUtils.sanitize_input(req.github, 200) if req.github else "",
            blog=SecurityUtils.sanitize_input(req.blog, 200) if req.blog else "",
            portfolio=SecurityUtils.sanitize_input(req.portfolio, 200) if req.portfolio else "",
            personalWebsite=SecurityUtils.sanitize_input(req.personalWebsite, 200) if req.personalWebsite else "",
            school=SecurityUtils.sanitize_input(req.school, 100),
            major=SecurityUtils.sanitize_input(req.major, 100),
            educationLevel=req.educationLevel,
            graduation=SecurityUtils.sanitize_input(req.graduation, 20),
            gpa=SecurityUtils.sanitize_input(req.gpa, 20) if req.gpa else "",
            courses=SecurityUtils.sanitize_input(req.courses, 500) if req.courses else "",
            skills=SecurityUtils.sanitize_input(req.skills, 1000) if req.skills else "",
            certificates=SecurityUtils.sanitize_input(req.certificates, 500) if req.certificates else "",
            selfEvaluation=SecurityUtils.sanitize_input(req.selfEvaluation, 1000) if req.selfEvaluation else "",
            jd=SecurityUtils.sanitize_input(req.jd, 5000),
            templateId=req.templateId,
            experiences=[
                ExperienceItem(
                    id=exp.id,
                    type=exp.type,
                    role=SecurityUtils.sanitize_input(exp.role, 100),
                    company=SecurityUtils.sanitize_input(exp.company, 100),
                    date=SecurityUtils.sanitize_input(exp.date, 50),
                    content=SecurityUtils.sanitize_input(exp.content, 5000)
                ) for exp in req.experiences
            ],
            starResults=req.starResults,
            suggestions=req.suggestions
        )
        
        html = await call_doubao_generate_resume(safe_req)
        
        # 净化生成的HTML
        safe_html = SecurityUtils.sanitize_generated_html(html)
        
        return GenerateResumeResponse(html=safe_html, success=True)
    except Exception as e:
        import logging
        logging.error(f"简历生成失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail="服务器内部错误，请稍后重试")

@app.post("/api/export/word")
async def export_word(req: ExportWordRequest):
    if not req.html or req.html.strip() == "":
        raise HTTPException(status_code=400, detail="简历内容不能为空")
    try:
        # 净化文件名，只允许安全字符
        raw_name = req.name.strip() or "简历"
        safe_name = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9._\-]', '_', raw_name)
        safe_name = safe_name[:50]  # 限制长度
        if not safe_name or safe_name.startswith('.'):
            safe_name = "简历"
        
        # 净化HTML内容
        safe_html = SecurityUtils.sanitize_generated_html(req.html)
        
        buffer = html_to_word(safe_html, safe_name, req.templateId)
        headers = {
            "Content-Disposition": f"attachment; filename*=UTF-8''{urllib.parse.quote(f'{safe_name}_简历.docx')}",
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "Access-Control-Expose-Headers": "Content-Disposition",
        }
        return StreamingResponse(buffer, headers=headers, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        import logging
        logging.error(f"Word文档导出失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail="文档导出失败，请稍后重试")

@app.post("/api/resume/upload")
async def upload_resume(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="未选择文件")
    
    ext = get_file_extension(file.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400, 
            detail=f"不支持的文件格式：{ext}。支持格式：PDF、DOC、DOCX、TXT"
        )
    
    content = await file.read()
    if len(content) == 0:
        raise HTTPException(status_code=400, detail="文件内容为空")
    
    if len(content) > MAX_FILE_SIZE:
        size_mb = len(content) / (1024 * 1024)
        raise HTTPException(
            status_code=400, 
            detail=f"文件大小({size_mb:.1f}MB)超过限制(15MB)"
        )
    
    file_hash = hashlib.sha256(content).hexdigest()
    
    safe_filename = generate_safe_filename(file.filename)
    file_path = os.path.join(UPLOAD_DIR, safe_filename)
    
    try:
        with open(file_path, "wb") as f:
            f.write(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")
    
    extracted_text = extract_text_from_file(file_path, ext)
    
    if not extracted_text or len(extracted_text.strip()) < 10:
        os.remove(file_path)
        raise HTTPException(status_code=400, detail="无法从文件中提取有效文本内容，请确保文件非扫描图片且包含可读文字")
    
    resume_id = uuid.uuid4().hex
    
    resume_store[resume_id] = {
        "id": resume_id,
        "original_filename": file.filename,
        "stored_filename": safe_filename,
        "file_path": file_path,
        "file_size": len(content),
        "file_hash": file_hash,
        "extension": ext,
        "extracted_text": extracted_text,
        "upload_time": datetime.now().isoformat(),
    }
    
    return {
        "success": True,
        "resume_id": resume_id,
        "filename": file.filename,
        "file_size": len(content),
        "extension": ext,
        "extracted_text_length": len(extracted_text),
        "message": "简历上传成功"
    }

@app.get("/api/resume/{resume_id}")
async def get_resume_info(resume_id: str):
    if resume_id not in resume_store:
        raise HTTPException(status_code=404, detail="简历不存在或已过期")
    info = resume_store[resume_id]
    return {
        "resume_id": info["id"],
        "filename": info["original_filename"],
        "file_size": info["file_size"],
        "extension": info["extension"],
        "upload_time": info["upload_time"],
        "text_length": len(info["extracted_text"]),
    }

@app.delete("/api/resume/{resume_id}")
async def delete_resume(resume_id: str):
    if resume_id not in resume_store:
        raise HTTPException(status_code=404, detail="简历不存在或已过期")
    info = resume_store.pop(resume_id)
    try:
        if os.path.exists(info["file_path"]):
            os.remove(info["file_path"])
    except:
        pass
    return {"success": True, "message": "简历已删除"}

@app.post("/api/resume/analyze")
async def analyze_resume(resume_id: str = Form(...)):
    if resume_id not in resume_store:
        raise HTTPException(status_code=404, detail="简历不存在或已过期")
    if not API_KEY:
        raise HTTPException(status_code=500, detail="API Key 未配置")

    resume_text = resume_store[resume_id]["extracted_text"]

    system_prompt = """你是一位拥有10年经验的资深HR总监和简历分析专家。
你的任务是对用户上传的简历进行全面、专业的分析。

【分析维度】
1. 整体评价：简历的整体质量和专业程度
2. 结构分析：简历结构是否完整，板块是否齐全
3. 内容分析：经历描述是否具体、量化，是否使用STAR法则
4. 技能匹配：技能描述是否与目标岗位匹配
5. 亮点与不足：明确指出简历的亮点和需要改进的地方
6. 优化建议：给出3-5条具体的优化建议

【输出格式】
必须输出严格的JSON格式，不要包含markdown标记：
{
    "overall_score": 75,
    "overall_comment": "整体评价...",
    "structure_analysis": "结构分析...",
    "content_analysis": "内容分析...",
    "highlights": ["亮点1", "亮点2"],
    "weaknesses": ["不足1", "不足2"],
    "suggestions": ["建议1", "建议2", "建议3"],
    "keywords": ["关键词1", "关键词2"]
}"""

    user_prompt = f"""请分析以下简历内容：

{resume_text[:4000]}

请给出全面、专业的分析，只输出JSON格式。"""

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                BASE_URL,
                headers={"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"},
                json={
                    "model": MODEL_ID,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    "temperature": 0.7,
                    "max_tokens": 2000
                }
            )
            response.raise_for_status()
            result = response.json()
            ai_content = clean_json_response(result["choices"][0]["message"]["content"])
            analysis = json.loads(ai_content)
            return {"success": True, "analysis": analysis}
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=502, detail=f"AI服务请求失败({e.response.status_code})")
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="AI分析超时，请稍后重试")
    except json.JSONDecodeError:
        return {"success": False, "analysis": {"overall_score":0,"overall_comment":"AI返回数据解析异常","structure_analysis":"","content_analysis":"","highlights":[],"weaknesses":["AI解析异常"],"suggestions":["请稍后重试"],"keywords":[]}}
    except Exception as e:
        import logging
        logging.error(f"Resume analysis error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"分析失败: {str(e)}")

@app.post("/api/resume/chat")
async def chat_with_resume(req: ResumeChatRequest):
    if req.resume_id not in resume_store:
        raise HTTPException(status_code=404, detail="简历不存在或已过期")
    if not API_KEY:
        raise HTTPException(status_code=500, detail="API Key 未配置")
    
    resume_text = resume_store[req.resume_id]["extracted_text"]
    
    # 净化用户消息
    safe_message = SecurityUtils.sanitize_input(req.message, 2000)
    if not safe_message.strip():
        raise HTTPException(status_code=400, detail="消息内容不能为空")
    
    system_prompt = f"""你是"星途简历"AI助手，一位专业的简历顾问和职业规划师。
你可以基于用户上传的简历内容，回答关于简历的各种问题，包括但不限于：

1. 简历内容解读和分析
2. 简历优化建议（措辞、结构、内容）
3. 职位匹配度评估
4. 面试准备建议
5. 职业发展规划建议
6. 技能提升建议

【用户简历内容】
{resume_text[:3000]}

【规则】
1. 回答必须基于简历实际内容，不得虚构信息
2. 给出具体、可操作的建议，避免空泛表述
3. 如果用户问题超出简历范围，可以适当拓展但需说明
4. 保持专业、友好的语气
5. 回答简洁明了，重点突出"""

    messages = [{"role": "system", "content": system_prompt}]
    
    if req.chat_history:
        for msg in req.chat_history[-10:]:
            role = msg.get("role", "user")
            content = SecurityUtils.sanitize_input(msg.get("content", ""), 2000)
            if role in ("user", "assistant") and content:
                messages.append({"role": role, "content": content})
    
    messages.append({"role": "user", "content": safe_message})

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                BASE_URL,
                headers={"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"},
                json={
                    "model": MODEL_ID,
                    "messages": messages,
                    "temperature": 0.7,
                    "max_tokens": 1500
                }
            )
            response.raise_for_status()
            result = response.json()
            ai_reply = result["choices"][0]["message"]["content"]

            safe_reply = SecurityUtils.sanitize_input(ai_reply, 5000)

            return {"success": True, "reply": safe_reply}
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=502, detail=f"AI服务请求失败({e.response.status_code})")
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="AI响应超时，请稍后重试")
    except (KeyError, IndexError) as e:
        import logging
        logging.error(f"Chat API format error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=502, detail="AI服务返回数据格式异常")
    except Exception as e:
        import logging
        logging.error(f"Chat error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"处理消息时出错: {str(e)}")

@app.get("/api/health")
async def health_check():
    return {"status": "ok", "model": MODEL_ID, "key_ok": bool(API_KEY)}

# 挂载前端页面（确保index.html在同一目录）
app.mount("/", StaticFiles(directory=".", html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)